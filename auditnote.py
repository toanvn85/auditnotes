import streamlit as st
st.set_page_config(page_title="Ứng dụng Đánh giá ISO", layout="wide")

import pandas as pd
import numpy as np
import gspread
import hashlib
import time
import os
import re
import io
import json
import base64
from datetime import datetime
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
import plotly.express as px
from PIL import Image
import pillow_heif
import requests

# Import libraries for PDF and Word export
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("Thư viện python-docx không khả dụng. Chức năng xuất Word sẽ bị hạn chế.")

# Predefined constants for ISO audit
ISO_CLAUSE_DATA = {
    "4": "Context of the organization",
    "4.1": "Understanding the organization and its context",
    "4.2": "Understanding the needs and expectations of interested parties",
    "4.3": "Determining the scope of the energy management system",
    "4.4": "Energy management system",
    "5": "Leadership",
    "5.1": "Leadership and commitment",
    "5.2": "Energy policy",
    "5.3": "Organization roles, responsibilities and authorities",
    "6": "Planning",
    "6.1": "Actions to address risks and opportunities",
    "6.2": "Objectives, energy targets and planning to achieve them",
    "6.3": "Energy review",
    "6.4": "Energy performance indicators",
    "6.5": "Energy baseline",
    "6.6": "Planning for collection of energy data",
    "7": "Support",
    "7.1": "Resources",
    "7.2": "Competence",
    "7.3": "Awareness",
    "7.4": "Communication",
    "7.5": "Documented information",
    "8": "Operation",
    "8.1": "Operational planning and control",
    "8.2": "Design",
    "8.3": "Procurement",
    "9": "Performance evaluation",
    "9.1": "Monitoring, measurement, analysis and evaluation of energy performance and the EnMS",
    "9.2": "Internal audit",
    "9.3": "Management review",
    "10": "Improvement",
    "10.1": "Nonconformity and corrective action",
    "10.2": "Continual improvement"
}

# Results definitions
AUDIT_RESULTS = {
    "NCA": "Phát hiện không phù hợp loại A",
    "NCB": "Phát hiện không phù hợp loại B",
    "PI": "Cơ hội cải tiến",
    "CM": "Phù hợp"
}

# ------------ Cấu hình logo 3×3 cm ~ 113×113 px ------------
LOGO_WIDTH, LOGO_HEIGHT = int(3/2.54*96), int(3/2.54*96)
def display_logos():
    """Tự động tìm và hiển thị logo1.*, logo2.* và logo3.* với đa định dạng."""
    c1, c2, c3, c4, c5 = st.columns(5)
    for col, base in ((c1, "logo1"), (c3, "logo2"), (c5, "logo3")):
        found = None
        for ext in ("png","jpg","jpeg","gif"):
            path = f"{base}.{ext}"
            if os.path.exists(path):
                found = path
                break
        if found:
            try:
                img = Image.open(found).resize((LOGO_WIDTH, LOGO_HEIGHT))
                col.image(img)
            except Exception as e:
                col.error(f"Lỗi đọc {found}: {e}")
        else:
            col.warning(f"Thiếu {base}.(png/jpg/jpeg/gif)")

# ------------ Thiết lập Google Sheets ------------
SCOPE = ["https://www.googleapis.com/auth/spreadsheets",
         "https://www.googleapis.com/auth/drive"]

def retry(func, tries=5, delay=1, mult=2):
    for i in range(tries):
        try:
            return func()
        except gspread.exceptions.APIError as e:
            if "429" in str(e) and i < tries-1:
                st.warning(f"Giới hạn tốc độ, thử lại sau {delay}s…")
                time.sleep(delay)
                delay *= mult
            else:
                raise

@st.cache_resource
def gclient():
    if os.path.exists("credentials.json"):
        creds = Credentials.from_service_account_file("credentials.json", scopes=SCOPE)
    else:
        creds = Credentials.from_service_account_info(
            st.secrets["gcp_service_account"], scopes=SCOPE
        )
    return gspread.authorize(creds)

def get_gdrive_service(_credentials):
    try:
        service = build('drive', 'v3', credentials=_credentials)
        return service
    except Exception as e:
        st.error(f"🔥 Lỗi kết nối Google Drive: {e}")
        return None

def ensure_header(ws, header):
    cur = [c.lower() for c in ws.row_values(1)]
    tgt = [h.lower() for h in header]
    if cur != tgt:
        ws.resize(rows=max(ws.row_count,1), cols=len(header))
        ws.update(f"A1:{chr(64+len(header))}1", [header])

@st.cache_resource(ttl=3600)
def gws():
    cli = gclient()
    
    # Auditors_DB
    try: adb = cli.open("Auditors_DB")
    except gspread.exceptions.SpreadsheetNotFound:
        adb = cli.create("Auditors_DB")
        adb.add_worksheet("Auditors", rows=10, cols=5)
    
    auditors_ws = adb.worksheet("Auditors")
    ensure_header(auditors_ws, [
        "fullname", "position", "email", "password", "last_login"
    ])
    
    # Default auditor nếu cần
    if len(auditors_ws.get_all_values()) == 1:
        pw0 = hashlib.sha256("auditor123".encode()).hexdigest()
        auditors_ws.append_row([
            "Đánh giá viên", "Trưởng đoàn", "auditor@example.com", pw0, ""
        ])
    
    # Audit_Notes
    try: 
        notes_wb = cli.open("Audit_Notes")
        notes_ws = notes_wb.worksheet("Notes")
    except gspread.exceptions.SpreadsheetNotFound:
        notes_wb = cli.create("Audit_Notes")
        notes_ws = notes_wb.sheet1
        notes_ws.update_title("Notes")
        ensure_header(notes_ws, [
            "company", "address", "department", "person", "audit_time",
            "frame_id", "panel_id", "clause", "clause_name", "requirements",
            "evidence", "image_url", "result", "auditor", "timestamp"
        ])
    except gspread.exceptions.WorksheetNotFound:
        notes_wb = cli.open("Audit_Notes")
        notes_ws = notes_wb.add_worksheet("Notes", rows=1, cols=15)
        ensure_header(notes_ws, [
            "company", "address", "department", "person", "audit_time",
            "frame_id", "panel_id", "clause", "clause_name", "requirements",
            "evidence", "image_url", "result", "auditor", "timestamp"
        ])
    
    # Audit_Participants
    try: 
        part_wb = cli.open("Audit_Participants") 
        part_ws = part_wb.worksheet("Participants")
    except gspread.exceptions.SpreadsheetNotFound:
        part_wb = cli.create("Audit_Participants")
        part_ws = part_wb.sheet1
        part_ws.update_title("Participants")
        ensure_header(part_ws, [
            "company", "frame_id", "fullname", "position", "role"
        ])
    except gspread.exceptions.WorksheetNotFound:
        part_wb = cli.open("Audit_Participants")
        part_ws = part_wb.add_worksheet("Participants", rows=1, cols=5)
        ensure_header(part_ws, [
            "company", "frame_id", "fullname", "position", "role"
        ])
    
    return {
        "auditors": auditors_ws,
        "notes_wb": notes_wb,
        "notes": notes_ws,
        "participants": part_ws
    }

# ------------ DataFrame Helpers ------------
def _df(ws):
    data = ws.get_all_values()
    if len(data) <= 1:
        cols = [c.lower() for c in data[0]] if data else []
        return pd.DataFrame(columns=cols)
    return pd.DataFrame(data[1:], columns=[c.lower() for c in data[0]])

@st.cache_data(ttl=300)
def df_auditors():   return _df(gws()["auditors"])
@st.cache_data(ttl=300)
def df_notes():      return _df(gws()["notes"])
@st.cache_data(ttl=300)
def df_participants(): return _df(gws()["participants"])

# ------------ Utilities ------------
hash_pw = lambda x: hashlib.sha256(x.encode()).hexdigest()
verify_pw = lambda s, p: s.strip() == hash_pw(p.strip())
sheet_name = lambda em: re.sub(r'[^A-Za-z0-9_-]', '_', em)[:100]

# --- Image Handling ---
def convert_heic_to_jpeg(file_object):
    try:
        file_object.seek(0)
        heif_file = pillow_heif.read_heif(file_object.read())
        image = Image.frombytes(
            heif_file.mode,
            heif_file.size,
            heif_file.data,
            "raw"
        )
        output = io.BytesIO()
        image.save(output, format="JPEG")
        output.seek(0)
        return output
    except Exception as e:
        st.error(f"❌ Không thể chuyển .heic -> .jpg: {e}")
        return None

def upload_image_to_drive(drive_service, file_object, folder_id):
    if not drive_service or not file_object or not folder_id:
        st.error("upload_image_to_drive: Đầu vào không hợp lệ")
        return None

    file_ext = file_object.name.lower().split('.')[-1]
    file_name_no_ext = file_object.name.rsplit('.', 1)[0]

    try:
        if file_ext in ['heic', 'heif']:
            converted_image = convert_heic_to_jpeg(file_object)
            if not converted_image:
                return None
            media = MediaIoBaseUpload(converted_image, mimetype='image/jpeg', resumable=True)
            new_filename = file_name_no_ext + ".jpg"
            file_metadata = {'name': new_filename, 'parents': [folder_id]}
        else:
            file_object.seek(0)
            media_content = io.BytesIO(file_object.getvalue())
            media = MediaIoBaseUpload(media_content, mimetype=file_object.type, resumable=True)
            file_metadata = {'name': file_object.name, 'parents': [folder_id]}

        file = drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()
        file_id = file.get('id')
        if not file_id:
            return None

        permission = {'type': 'anyone', 'role': 'reader'}
        drive_service.permissions().create(fileId=file_id, body=permission).execute()

        # Trả về URL trực tiếp của ảnh
        image_url = f"https://drive.google.com/uc?export=view&id={file_id}"
        return image_url

    except Exception as e:
        st.error(f"Lỗi khi upload ảnh: {e}")
        return None

# --- CSS Styling ---
def load_css():
    st.markdown("""
    <style>
    .block-container {
        padding: 2rem;
    }
    .stTextInput>div>input, .stTextArea>div>textarea {
        border-radius: 0.5rem;
        border: 1px solid #D3D3D3;
    }
    .stButton>button {
        border-radius: 8px;
        padding: 8px 20px;
    }
    .download-button {
        display: inline-block;
        padding: 0.5em 1em;
        text-decoration: none;
        color: white;
        background-color: #0066cc;
        border-radius: 5px;
        font-weight: bold;
        margin: 0.5em 0;
        text-align: center;
    }
    .download-button:hover {
        background-color: #0052a3;
    }
    </style>
    """, unsafe_allow_html=True)

# ============ Trang Đăng nhập ============
def page_login():
    display_logos()
    st.title("Đăng nhập Hệ thống Đánh giá ISO")
    
    email = st.text_input("Email đánh giá viên")
    password = st.text_input("Mật khẩu", type="password")
    
    col1, col2 = st.columns(2)
    
    if col1.button("Đăng nhập"):
        auditors = df_auditors()
        if email == "admin" and password == "admin123":
            # Admin login for testing
            st.session_state.user = {
                "email": "admin",
                "fullname": "Admin",
                "position": "Administrator"
            }
            st.session_state.is_logged_in = True
            st.rerun()
        elif not auditors.empty:
            user = auditors[auditors['email'] == email]
            if not user.empty:
                stored_pw = user.iloc[0]['password']
                if verify_pw(stored_pw, password):
                    st.session_state.user = {
                        "email": email,
                        "fullname": user.iloc[0]['fullname'],
                        "position": user.iloc[0]['position']
                    }
                    
                    # Cập nhật thời gian đăng nhập
                    auditor_idx = auditors[auditors['email'] == email].index[0] + 2
                    gws()["auditors"].update_cell(auditor_idx, 5, datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
                    df_auditors.clear()
                    
                    st.session_state.is_logged_in = True
                    st.rerun()
                else:
                    st.error("Mật khẩu không đúng!")
            else:
                st.error("Không tìm thấy email đánh giá viên!")
    
    if col2.button("Đăng ký"):
        st.session_state.show_register = True
        st.rerun()
    
    # Form đăng ký
    if st.session_state.get("show_register", False):
        st.subheader("Đăng ký đánh giá viên mới")
        with st.form("register_form"):
            fullname = st.text_input("Họ và tên")
            position = st.text_input("Chức vụ")
            reg_email = st.text_input("Email")
            reg_password = st.text_input("Mật khẩu", type="password")
            confirm_password = st.text_input("Xác nhận mật khẩu", type="password")
            
            submit = st.form_submit_button("Đăng ký")
            
            if submit:
                if reg_password != confirm_password:
                    st.error("Mật khẩu không khớp!")
                elif not fullname or not position or not reg_email:
                    st.error("Vui lòng điền đầy đủ thông tin!")
                else:
                    auditors = df_auditors()
                    if not auditors.empty and (auditors['email'] == reg_email).any():
                        st.error("Email đã tồn tại!")
                    else:
                        hashed_pw = hash_pw(reg_password)
                        gws()["auditors"].append_row([
                            fullname, position, reg_email, hashed_pw, datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        ])
                        df_auditors.clear()
                        st.success("Đăng ký thành công! Vui lòng đăng nhập.")
                        st.session_state.show_register = False
                        st.rerun()

# ============ Trang Đổi Mật Khẩu ============
def page_change_password():
    display_logos()
    st.title("Đổi mật khẩu")
    
    with st.form("change_password_form"):
        current_pw = st.text_input("Mật khẩu hiện tại", type="password")
        new_pw = st.text_input("Mật khẩu mới", type="password")
        confirm_pw = st.text_input("Xác nhận mật khẩu mới", type="password")
        
        submit = st.form_submit_button("Đổi mật khẩu")
        
        if submit:
            if new_pw != confirm_pw:
                st.error("Mật khẩu mới không khớp!")
            else:
                auditors = df_auditors()
                user = auditors[auditors['email'] == st.session_state.user["email"]]
                if not user.empty:
                    stored_pw = user.iloc[0]['password']
                    if verify_pw(stored_pw, current_pw):
                        # Cập nhật mật khẩu mới
                        hashed_pw = hash_pw(new_pw)
                        auditor_idx = auditors[auditors['email'] == st.session_state.user["email"]].index[0] + 2
                        gws()["auditors"].update_cell(auditor_idx, 4, hashed_pw)
                        df_auditors.clear()
                        st.success("Đổi mật khẩu thành công!")
                    else:
                        st.error("Mật khẩu hiện tại không đúng!")
    
    if st.button("Quay lại"):
        st.session_state.page = "main"
        st.rerun()

# --- Export Functions ---
def get_download_link(file_content, file_name, display_text):
    """Tạo download link cho các file được tạo."""
    b64 = base64.b64encode(file_content).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{file_name}" class="download-button">{display_text}</a>'
    return href

def process_image_for_export(image_url):
    """Xử lý URL hình ảnh để sử dụng trong export PDF và Word"""
    if not image_url:
        return None
    
    try:
        response = requests.get(image_url)
        img = Image.open(io.BytesIO(response.content))
        return img
    except Exception as e:
        st.error(f"Lỗi xử lý ảnh: {e}")
        return None
# ============ Export Functions ============
def export_to_pdf(company_name, audit_data, participants_data):
    """Tạo file PDF từ dữ liệu audit."""
    buffer = io.BytesIO()
    
    # Cố gắng đăng ký font hỗ trợ tiếng Việt nếu có
    try:
        pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
        font_name = 'DejaVuSans'
    except:
        font_name = 'Helvetica'
    
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4))
    styles = getSampleStyleSheet()
    
    # Tạo style cho tiêu đề và nội dung
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontName=font_name,
        fontSize=16,
        alignment=1,
        spaceAfter=12
    )
    
    subtitle_style = ParagraphStyle(
        'Subtitle',
        parent=styles['Heading2'],
        fontName=font_name,
        fontSize=14,
        alignment=1,
        spaceAfter=10
    )
    
    normal_style = ParagraphStyle(
        'Normal',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        spaceAfter=6
    )
    
    content = []
    
    # Tiêu đề
    content.append(Paragraph(f"BÁO CÁO ĐÁNH GIÁ ISO", title_style))
    content.append(Paragraph(f"Công ty: {company_name}", subtitle_style))
    content.append(Spacer(1, 10))
    
    # Thông tin chung
    if audit_data and len(audit_data) > 0:
        general_info = [
            ["Bộ phận được đánh giá:", audit_data[0]['department']],
            ["Người đối ứng:", audit_data[0]['person']],
            ["Thời gian đánh giá:", audit_data[0]['audit_time']],
            ["Địa chỉ:", audit_data[0]['address']]
        ]
        
        t = Table(general_info, colWidths=[150, 400])
        t.setStyle(TableStyle([
            ('FONT', (0, 0), (-1, -1), font_name, 10),
            ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ]))
        content.append(t)
        content.append(Spacer(1, 10))
    
    # Danh sách thành viên tham gia và đánh giá viên
    if participants_data:
        content.append(Paragraph("THÀNH VIÊN THAM GIA", subtitle_style))
        
        company_participants = [p for p in participants_data if p['role'] == 'company']
        if company_participants:
            participant_data = [["Họ và tên", "Chức vụ"]]
            for p in company_participants:
                participant_data.append([p['fullname'], p['position']])
            
            t = Table(participant_data, colWidths=[275, 275])
            t.setStyle(TableStyle([
                ('FONT', (0, 0), (-1, -1), font_name, 10),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ]))
            content.append(t)
        
        content.append(Spacer(1, 10))
        content.append(Paragraph("ĐÁNH GIÁ VIÊN", subtitle_style))
        
        auditor_participants = [p for p in participants_data if p['role'] == 'auditor']
        if auditor_participants:
            auditor_data = [["Họ và tên", "Chức vụ"]]
            for p in auditor_participants:
                auditor_data.append([p['fullname'], p['position']])
            
            t = Table(auditor_data, colWidths=[275, 275])
            t.setStyle(TableStyle([
                ('FONT', (0, 0), (-1, -1), font_name, 10),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ]))
            content.append(t)
        
        content.append(Spacer(1, 15))
    
    # Phân tích theo Frame
    if audit_data:
        frames = {}
        for item in audit_data:
            frame_id = item['frame_id']
            if frame_id not in frames:
                frames[frame_id] = []
            frames[frame_id].append(item)
        
        for frame_id, frame_items in frames.items():
            content.append(Paragraph(f"FRAME {frame_id}", subtitle_style))
            
            # Thống kê kết quả
            results = {'NCA': 0, 'NCB': 0, 'PI': 0, 'CM': 0}
            for item in frame_items:
                if item['result'] in results:
                    results[item['result']] += 1
            
            result_data = [["NCA", "NCB", "PI", "CM"]]
            result_data.append([str(results['NCA']), str(results['NCB']), str(results['PI']), str(results['CM'])])
            
            t = Table(result_data, colWidths=[137.5, 137.5, 137.5, 137.5])
            t.setStyle(TableStyle([
                ('FONT', (0, 0), (-1, -1), font_name, 10),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ]))
            content.append(t)
            content.append(Spacer(1, 10))
            
            # Dữ liệu chi tiết
            for item in frame_items:
                data = [
                    ["Điều khoản", "Tên điều khoản", "Các yêu cầu Tiêu chuẩn/Chuẩn mực đánh giá", "Bằng chứng đánh giá", "Kết quả đánh giá"],
                    [item['clause'], item['clause_name'], item['requirements'], item['evidence'], item['result']]
                ]
                
                t = Table(data, colWidths=[70, 100, 150, 150, 80])
                t.setStyle(TableStyle([
                    ('FONT', (0, 0), (-1, -1), font_name, 9),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ]))
                content.append(t)
                content.append(Spacer(1, 5))
                
                # Thêm hình ảnh nếu có
                if item['image_url']:
                    try:
                        img = process_image_for_export(item['image_url'])
                        if img:
                            img_data = io.BytesIO()
                            img.save(img_data, format='JPEG')
                            img_data.seek(0)
                            
                            # Resize image to fit in PDF
                            width, height = img.size
                            aspect = width / height
                            if width > 300:
                                width = 300
                                height = width / aspect
                            
                            content.append(Paragraph(f"Hình ảnh bằng chứng:", normal_style))
                            content.append(RLImage(img_data, width=width, height=height))
                            content.append(Spacer(1, 10))
                    except Exception as e:
                        content.append(Paragraph(f"[Không thể hiển thị hình ảnh: {e}]", normal_style))
            
            content.append(Spacer(1, 20))
    
    # Thêm ngày xuất báo cáo
    current_date = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    content.append(Spacer(1, 10))
    content.append(Paragraph(f"Báo cáo được xuất ngày: {current_date}", normal_style))
    
    doc.build(content)
    pdf = buffer.getvalue()
    buffer.close()
    return pdf

def export_to_word(company_name, audit_data, participants_data):
    """Tạo file Word từ dữ liệu audit."""
    doc = Document()
    
    # Thiết lập font và cỡ chữ mặc định
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    
    # Tiêu đề
    heading = doc.add_heading('BÁO CÁO ĐÁNH GIÁ ISO', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    company_heading = doc.add_heading(f'Công ty: {company_name}', level=2)
    company_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thông tin chung
    if audit_data and len(audit_data) > 0:
        doc.add_heading('Thông tin chung', level=3)
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        cells = table.rows[0].cells
        cells[0].text = 'Bộ phận được đánh giá:'
        cells[1].text = audit_data[0]['department']
        
        cells = table.rows[1].cells
        cells[0].text = 'Người đối ứng:'
        cells[1].text = audit_data[0]['person']
        
        cells = table.rows[2].cells
        cells[0].text = 'Thời gian đánh giá:'
        cells[1].text = audit_data[0]['audit_time']
        
        cells = table.rows[3].cells
        cells[0].text = 'Địa chỉ:'
        cells[1].text = audit_data[0]['address']
    
    # Danh sách thành viên tham gia và đánh giá viên
    if participants_data:
        doc.add_heading('THÀNH VIÊN THAM GIA', level=3)
        company_participants = [p for p in participants_data if p['role'] == 'company']
        if company_participants:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Họ và tên'
            hdr_cells[1].text = 'Chức vụ'
            
            for p in company_participants:
                row_cells = table.add_row().cells
                row_cells[0].text = p['fullname']
                row_cells[1].text = p['position']
        
        doc.add_heading('ĐÁNH GIÁ VIÊN', level=3)
        auditor_participants = [p for p in participants_data if p['role'] == 'auditor']
        if auditor_participants:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Họ và tên'
            hdr_cells[1].text = 'Chức vụ'
            
            for p in auditor_participants:
                row_cells = table.add_row().cells
                row_cells[0].text = p['fullname']
                row_cells[1].text = p['position']
    
    # Phân tích theo Frame
    if audit_data:
        frames = {}
        for item in audit_data:
            frame_id = item['frame_id']
            if frame_id not in frames:
                frames[frame_id] = []
            frames[frame_id].append(item)
        
        for frame_id, frame_items in frames.items():
            doc.add_heading(f'FRAME {frame_id}', level=3)
            
            # Thống kê kết quả
            results = {'NCA': 0, 'NCB': 0, 'PI': 0, 'CM': 0}
            for item in frame_items:
                if item['result'] in results:
                    results[item['result']] += 1
            
            table = doc.add_table(rows=2, cols=4)
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'NCA'
            hdr_cells[1].text = 'NCB'
            hdr_cells[2].text = 'PI'
            hdr_cells[3].text = 'CM'
            
            result_cells = table.rows[1].cells
            result_cells[0].text = str(results['NCA'])
            result_cells[1].text = str(results['NCB'])
            result_cells[2].text = str(results['PI'])
            result_cells[3].text = str(results['CM'])
            
            doc.add_paragraph()
            
            # Dữ liệu chi tiết
            for idx, item in enumerate(frame_items):
                doc.add_paragraph(f"Điều mục {idx+1}:")
                
                table = doc.add_table(rows=2, cols=5)
                table.style = 'Table Grid'
                
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Điều khoản'
                hdr_cells[1].text = 'Tên điều khoản'
                hdr_cells[2].text = 'Các yêu cầu Tiêu chuẩn/Chuẩn mực đánh giá'
                hdr_cells[3].text = 'Bằng chứng đánh giá'
                hdr_cells[4].text = 'Kết quả đánh giá'
                
                data_cells = table.rows[1].cells
                data_cells[0].text = item['clause']
                data_cells[1].text = item['clause_name']
                data_cells[2].text = item['requirements']
                data_cells[3].text = item['evidence']
                data_cells[4].text = item['result']
                
                # Thêm hình ảnh nếu có
                if item['image_url']:
                    try:
                        doc.add_paragraph("Hình ảnh bằng chứng:")
                        img = process_image_for_export(item['image_url'])
                        if img:
                            with io.BytesIO() as img_stream:
                                img.save(img_stream, format='JPEG')
                                img_stream.seek(0)
                                doc.add_picture(img_stream, width=Inches(4))
                    except Exception as e:
                        doc.add_paragraph(f"[Không thể hiển thị hình ảnh: {e}]")
                
                doc.add_paragraph()
    
    # Thêm ngày xuất báo cáo
    current_date = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    doc.add_paragraph(f"Báo cáo được xuất ngày: {current_date}")
    
    # Lưu vào memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    docx = buffer.getvalue()
    buffer.close()
    return docx
# ============ Trang Chính ============
def page_main():
    display_logos()
    st.title("Hệ thống Đánh giá ISO")
    
    # Sidebar
    with st.sidebar:
        st.subheader(f"Xin chào, {st.session_state.user['fullname']}")
        st.write(f"Chức vụ: {st.session_state.user['position']}")
        
        if st.button("Đổi mật khẩu"):
            st.session_state.page = "change_password"
            st.rerun()
        
        if st.button("Đăng xuất"):
            st.session_state.clear()
            st.rerun()
    
    # Initialize session state for audit data if needed
    if "audit_frames" not in st.session_state:
        st.session_state.audit_frames = {}
    
    if "current_frame" not in st.session_state:
        st.session_state.current_frame = "1"
    
    if "company_info" not in st.session_state:
        st.session_state.company_info = {
            "company_name": "",
            "address": "",
            "participants": [],
            "auditors": []
        }
    
    # Tab navigation
    tab1, tab2, tab3 = st.tabs(["Ghi chép đánh giá", "Xem lại đánh giá", "Xuất báo cáo"])
    
    with tab1:
        page_audit_entry()
    
    with tab2:
        page_audit_review()
    
    with tab3:
        page_export()

# ============ Trang Chính ============
def page_main():
    display_logos()
    st.title("Hệ thống Đánh giá ISO")
    
    # Sidebar
    with st.sidebar:
        st.subheader(f"Xin chào, {st.session_state.user['fullname']}")
        st.write(f"Chức vụ: {st.session_state.user['position']}")
        
        if st.button("Đổi mật khẩu"):
            st.session_state.page = "change_password"
            st.rerun()
        
        if st.button("Đăng xuất"):
            st.session_state.clear()
            st.rerun()
    
    # Initialize session state for audit data if needed
    if "audit_frames" not in st.session_state:
        st.session_state.audit_frames = {}
    
    if "current_frame" not in st.session_state:
        st.session_state.current_frame = "1"
    
    if "company_info" not in st.session_state:
        st.session_state.company_info = {
            "company_name": "",
            "address": "",
            "participants": [],
            "auditors": []
        }
    
    # Tab navigation
    tab1, tab2, tab3 = st.tabs(["Ghi chép đánh giá", "Xem lại đánh giá", "Xuất báo cáo"])
    
    with tab1:
        page_audit_entry()
    
    with tab2:
        page_audit_review()
    
    with tab3:
        page_export()

# ============ Trang Nhập liệu đánh giá ============
def page_audit_entry():
    # Display the company information form at the top
    with st.expander("Thông tin công ty", expanded=True):
        col1, col2 = st.columns(2)
        
        with col1:
            company_name = st.text_input("Tên công ty", 
                                         value=st.session_state.company_info["company_name"])
        
        with col2:
            address = st.text_input("Địa chỉ", 
                                    value=st.session_state.company_info["address"])
        
        # Save the company info
        st.session_state.company_info["company_name"] = company_name
        st.session_state.company_info["address"] = address
        
        # Participants management (company representatives)
        st.subheader("Thành phần tham gia họp khai mạc")
        
        # Initialize the participants list if empty
        if len(st.session_state.company_info["participants"]) == 0:
            st.session_state.company_info["participants"] = [{"fullname": "", "position": ""}]
        
        # Display existing participants
        participants_to_remove = []
        for i, participant in enumerate(st.session_state.company_info["participants"]):
            cols = st.columns([3, 3, 1])
            with cols[0]:
                st.session_state.company_info["participants"][i]["fullname"] = st.text_input(
                    f"Họ tên #{i+1}", 
                    value=participant["fullname"],
                    key=f"part_name_{i}"
                )
            with cols[1]:
                st.session_state.company_info["participants"][i]["position"] = st.text_input(
                    f"Chức vụ #{i+1}", 
                    value=participant["position"],
                    key=f"part_pos_{i}"
                )
            with cols[2]:
                if st.button("Xóa", key=f"del_part_{i}"):
                    participants_to_remove.append(i)
        
        # Remove marked participants
        for idx in sorted(participants_to_remove, reverse=True):
            st.session_state.company_info["participants"].pop(idx)
        
        # Add new participant button
        if st.button("➕ Thêm người tham gia"):
            st.session_state.company_info["participants"].append({"fullname": "", "position": ""})
            st.rerun()
        
        # Auditors management
        st.subheader("Đánh giá viên")
        
        # Initialize the auditors list if empty
        if len(st.session_state.company_info["auditors"]) == 0:
            # Add the current user as the first auditor
            st.session_state.company_info["auditors"] = [{
                "fullname": st.session_state.user["fullname"],
                "position": st.session_state.user["position"]
            }]
        
        # Display existing auditors
        auditors_to_remove = []
        for i, auditor in enumerate(st.session_state.company_info["auditors"]):
            cols = st.columns([3, 3, 1])
            with cols[0]:
                st.session_state.company_info["auditors"][i]["fullname"] = st.text_input(
                    f"Họ tên đánh giá viên #{i+1}", 
                    value=auditor["fullname"],
                    key=f"auditor_name_{i}"
                )
            with cols[1]:
                st.session_state.company_info["auditors"][i]["position"] = st.text_input(
                    f"Chức vụ đánh giá viên #{i+1}", 
                    value=auditor["position"],
                    key=f"auditor_pos_{i}"
                )
            with cols[2]:
                if i > 0:  # Don't allow removing the current user
                    if st.button("Xóa", key=f"del_auditor_{i}"):
                        auditors_to_remove.append(i)
        
        # Remove marked auditors
        for idx in sorted(auditors_to_remove, reverse=True):
            st.session_state.company_info["auditors"].pop(idx)
        
        # Add new auditor button
        if st.button("➕ Thêm đánh giá viên"):
            st.session_state.company_info["auditors"].append({"fullname": "", "position": ""})
            st.rerun()
    
    # Frame management
    st.subheader("Khung đánh giá")
    
    # Frame selector
    available_frames = list(st.session_state.audit_frames.keys())
    if not available_frames:
        available_frames = ["1"]
        st.session_state.audit_frames["1"] = {
            "department": "",
            "person": "",
            "audit_time": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "panels": {}
        }
    
    frame_cols = st.columns([2, 1])
    with frame_cols[0]:
        selected_frame = st.selectbox(
            "Chọn khung đánh giá", 
            options=available_frames,
            index=available_frames.index(st.session_state.current_frame)
        )
    
    with frame_cols[1]:
        if st.button("➕ Thêm khung đánh giá mới"):
            new_frame_id = str(len(available_frames) + 1)
            st.session_state.audit_frames[new_frame_id] = {
                "department": "",
                "person": "",
                "audit_time": datetime.now().strftime("%Y-%m-%d %H:%M"),
                "panels": {}
            }
            st.session_state.current_frame = new_frame_id
            st.rerun()
    
    st.session_state.current_frame = selected_frame
    current_frame_data = st.session_state.audit_frames[selected_frame]
    
    # Frame details
    with st.expander(f"Chi tiết khung đánh giá #{selected_frame}", expanded=True):
        department = st.text_input(
            "Bộ phận được đánh giá",
            value=current_frame_data.get("department", "")
        )
        st.session_state.audit_frames[selected_frame]["department"] = department
        
        person = st.text_input(
            "Người đối ứng",
            value=current_frame_data.get("person", "")
        )
        st.session_state.audit_frames[selected_frame]["person"] = person
        
        audit_time = st.text_input(
            "Thời gian đánh giá",
            value=current_frame_data.get("audit_time", datetime.now().strftime("%Y-%m-%d %H:%M"))
        )
        st.session_state.audit_frames[selected_frame]["audit_time"] = audit_time
    
    # Panel management for the current frame
    handle_panel_management(selected_frame)

# ============ Panel Management ============
def handle_panel_management(frame_id):
    """Handle panel management for the given frame ID"""
    current_frame = st.session_state.audit_frames[frame_id]
    
    # Initialize panels if empty
    if "panels" not in current_frame:
        current_frame["panels"] = {}
    
    # If no panels exist, create the first one
    if not current_frame["panels"]:
        current_frame["panels"]["1"] = {
            "items": []
        }
    
    # Panel selector
    st.subheader("Panel Notes")
    available_panels = list(current_frame["panels"].keys())
    
    panel_cols = st.columns([2, 1])
    with panel_cols[0]:
        selected_panel = st.selectbox(
            "Chọn Panel",
            options=available_panels,
            key=f"panel_select_{frame_id}"
        )
    
    with panel_cols[1]:
        if st.button("➕ Thêm Panel mới"):
            new_panel_id = str(len(available_panels) + 1)
            current_frame["panels"][new_panel_id] = {
                "items": []
            }
            st.rerun()
    
    # Display the selected panel's items
    handle_panel_items(frame_id, selected_panel)
    
    # Add new item to the panel
    with st.expander("Thêm mục đánh giá mới", expanded=True):
        with st.form(key=f"new_item_form_{frame_id}_{selected_panel}"):
            cols1 = st.columns(3)
            with cols1[0]:
                new_clause = st.text_input("Điều khoản", key=f"new_clause_{frame_id}_{selected_panel}")
            with cols1[1]:
                new_clause_name = st.text_input("Tên điều khoản", key=f"new_clause_name_{frame_id}_{selected_panel}")
            with cols1[2]:
                new_requirements = st.text_area("Các yêu cầu Tiêu chuẩn/Chuẩn mực đánh giá", 
                                              key=f"new_requirements_{frame_id}_{selected_panel}")
            
            cols2 = st.columns(2)
            with cols2[0]:
                new_evidence = st.text_area("Bằng chứng đánh giá", 
                                          key=f"new_evidence_{frame_id}_{selected_panel}")
            with cols2[1]:
                uploaded_file = st.file_uploader(
                    "Hình ảnh bằng chứng", 
                    type=["png", "jpg", "jpeg", "heic", "heif", "bmp"],
                    key=f"new_image_{frame_id}_{selected_panel}"
                )
            
            new_result = st.selectbox(
                "Kết quả đánh giá",
                options=["NCA", "NCB", "PI", "CM"],
                key=f"new_result_{frame_id}_{selected_panel}"
            )
            
            submit_button = st.form_submit_button("Thêm mục đánh giá")
            
            if submit_button:
                # Handle file upload if present
                image_url = None
                if uploaded_file:
                    creds = None
                    if os.path.exists("credentials.json"):
                        creds = Credentials.from_service_account_file("credentials.json", scopes=SCOPE)
                    else:
                        creds = Credentials.from_service_account_info(
                            st.secrets["gcp_service_account"], scopes=SCOPE
                        )
                    
                    drive_service = get_gdrive_service(creds)
                    folder_id = st.secrets["google_drive"]["folder_id"]
                    
                    with st.spinner("Đang tải ảnh lên..."):
                        image_url = upload_image_to_drive(drive_service, uploaded_file, folder_id)
                
                # Add the new item to the panel
                new_item = {
                    "clause": new_clause,
                    "clause_name": new_clause_name,
                    "requirements": new_requirements,
                    "evidence": new_evidence,
                    "image_url": image_url,
                    "result": new_result,
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                }
                
                current_frame["panels"][selected_panel]["items"].append(new_item)
                
                # Save to Google Sheets
                save_item_to_sheets(
                    st.session_state.company_info["company_name"],
                    st.session_state.company_info["address"],
                    current_frame["department"],
                    current_frame["person"],
                    current_frame["audit_time"],
                    frame_id,
                    selected_panel,
                    new_item,
                    st.session_state.user["email"]
                )
                
                st.success("Đã thêm mục đánh giá mới!")
                st.rerun()
    
    # Display statistics for this panel
    display_panel_statistics(current_frame["panels"][selected_panel]["items"])

def handle_panel_items(frame_id, panel_id):
    """Display and manage items in a panel"""
    current_frame = st.session_state.audit_frames[frame_id]
    current_panel = current_frame["panels"][panel_id]
    
    if not current_panel["items"]:
        st.info(f"Panel #{panel_id} chưa có mục đánh giá nào. Vui lòng thêm mục đánh giá mới.")
        return
    
    for idx, item in enumerate(current_panel["items"]):
        with st.expander(f"Mục đánh giá #{idx+1}: {item['clause']} - {item['clause_name']}", expanded=False):
            cols = st.columns(3)
            cols[0].write(f"**Điều khoản:** {item['clause']}")
            cols[1].write(f"**Tên điều khoản:** {item['clause_name']}")
            cols[2].write(f"**Kết quả đánh giá:** {item['result']}")
            
            st.write(f"**Các yêu cầu Tiêu chuẩn/Chuẩn mực đánh giá:**")
            st.text_area("", value=item['requirements'], disabled=True, key=f"req_{frame_id}_{panel_id}_{idx}")
            
            st.write(f"**Bằng chứng đánh giá:**")
            st.text_area("", value=item['evidence'], disabled=True, key=f"evi_{frame_id}_{panel_id}_{idx}")
            
            if item['image_url']:
                st.write("**Hình ảnh bằng chứng:**")
                st.image(item['image_url'])
            
            st.write(f"*Thời gian ghi nhận: {item['timestamp']}*")
            
            # Delete button
            if st.button("Xóa mục này", key=f"del_{frame_id}_{panel_id}_{idx}"):
                current_panel["items"].pop(idx)
                st.success("Đã xóa mục đánh giá!")
                st.rerun()

def display_panel_statistics(items):
    """Display statistics for panel items"""
    st.subheader("Thống kê kết quả đánh giá")
    
    results = {
        "NCA": 0,
        "NCB": 0,
        "PI": 0,
        "CM": 0
    }
    
    for item in items:
        result = item["result"]
        if result in results:
            results[result] += 1
    
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("NCA", results["NCA"])
    col2.metric("NCB", results["NCB"])
    col3.metric("PI", results["PI"])
    col4.metric("CM", results["CM"])
    
    return results

def save_item_to_sheets(company, address, department, person, audit_time, 
                       frame_id, panel_id, item, auditor_email):
    """Save an audit item to Google Sheets"""
    notes_ws = gws()["notes"]
    
    row = [
        company,
        address,
        department,
        person,
        audit_time,
        frame_id,
        panel_id,
        item["clause"],
        item["clause_name"],
        item["requirements"],
        item["evidence"],
        item["image_url"] if item["image_url"] else "",
        item["result"],
        auditor_email,
        item["timestamp"]
    ]
    
    notes_ws.append_row(row)
    
    # Also save participants if not already saved
    save_participants_to_sheets(company, frame_id)

def save_participants_to_sheets(company, frame_id):
    """Save participants to Google Sheets"""
    part_ws = gws()["participants"]
    
    # First, check if participants for this company and frame already exist
    part_df = df_participants()
    company_frame_parts = part_df[
        (part_df['company'] == company) & 
        (part_df['frame_id'] == frame_id)
    ]
    
    if not company_frame_parts.empty:
        return  # Already saved
    
    # Add company participants
    for participant in st.session_state.company_info["participants"]:
        if participant["fullname"] and participant["position"]:
            part_ws.append_row([
                company,
                frame_id,
                participant["fullname"],
                participant["position"],
                "company"
            ])
    
    # Add auditors
    for auditor in st.session_state.company_info["auditors"]:
        if auditor["fullname"] and auditor["position"]:
            part_ws.append_row([
                company,
                frame_id,
                auditor["fullname"],
                auditor["position"],
                "auditor"
            ])
    
    # Clear cache
    df_participants.clear()

# ============ Review Audit Data ============
def page_audit_review():
    """Page for reviewing past audit data"""
    st.subheader("Xem lại đánh giá")
    
    # Get all audit data from sheets
    notes_df = df_notes()
    
    if notes_df.empty:
        st.info("Chưa có dữ liệu đánh giá nào.")
        return
    
    # Get unique companies
    companies = notes_df['company'].unique()
    selected_company = st.selectbox("Chọn công ty", options=companies)
    
    # Filter by company
    company_data = notes_df[notes_df['company'] == selected_company]
    
    # Get frames for this company
    frames = company_data['frame_id'].unique()
    selected_frame = st.selectbox("Chọn khung đánh giá", options=frames)
    
    # Filter by frame
    frame_data = company_data[company_data['frame_id'] == selected_frame]
    
    # Display frame info
    if not frame_data.empty:
        first_row = frame_data.iloc[0]
        
        st.write(f"**Bộ phận được đánh giá:** {first_row['department']}")
        st.write(f"**Người đối ứng:** {first_row['person']}")
        st.write(f"**Thời gian đánh giá:** {first_row['audit_time']}")
        st.write(f"**Địa chỉ:** {first_row['address']}")
        
        # Get panels for this frame
        panels = frame_data['panel_id'].unique()
        
        for panel in panels:
            st.subheader(f"Panel #{panel}")
            
            # Get items for this panel
            panel_data = frame_data[frame_data['panel_id'] == panel]
            
            # Display panel statistics
            results = {
                "NCA": sum(panel_data['result'] == "NCA"),
                "NCB": sum(panel_data['result'] == "NCB"),
                "PI": sum(panel_data['result'] == "PI"),
                "CM": sum(panel_data['result'] == "CM")
            }
            
            col1, col2, col3, col4 = st.columns(4)
            col1.metric("NCA", results["NCA"])
            col2.metric("NCB", results["NCB"])
            col3.metric("PI", results["PI"])
            col4.metric("CM", results["CM"])
            
            # Display items
            for idx, item in panel_data.iterrows():
                with st.expander(f"Mục đánh giá: {item['clause']} - {item['clause_name']}", expanded=False):
                    cols = st.columns(3)
                    cols[0].write(f"**Điều khoản:** {item['clause']}")
                    cols[1].write(f"**Tên điều khoản:** {item['clause_name']}")
                    cols[2].write(f"**Kết quả đánh giá:** {item['result']}")
                    
                    st.write(f"**Các yêu cầu Tiêu chuẩn/Chuẩn mực đánh giá:**")
                    st.text_area("", value=item['requirements'], disabled=True, key=f"rev_req_{idx}")
                    
                    st.write(f"**Bằng chứng đánh giá:**")
                    st.text_area("", value=item['evidence'], disabled=True, key=f"rev_evi_{idx}")
                    
                    if item['image_url']:
                        st.write("**Hình ảnh bằng chứng:**")
                        st.image(item['image_url'])
                    
                    st.write(f"*Đánh giá bởi: {item['auditor']}*")
                    st.write(f"*Thời gian ghi nhận: {item['timestamp']}*")

# ============ Export Page ============
def page_export():
    """Page for exporting audit data"""
    st.subheader("Xuất báo cáo đánh giá")
    
    # Get audit data
    notes_df = df_notes()
    participants_df = df_participants()
    
    if notes_df.empty:
        st.info("Chưa có dữ liệu đánh giá nào để xuất báo cáo.")
        return
    
    # Get unique companies
    companies = notes_df['company'].unique()
    selected_company = st.selectbox("Chọn công ty", options=companies, key="export_company")
    
    # Filter by company
    company_data = notes_df[notes_df['company'] == selected_company]
    company_participants = participants_df[participants_df['company'] == selected_company]
    
    # Get frames for this company
    frames = company_data['frame_id'].unique()
    
    if len(frames) > 0:
        frame_option = st.radio(
            "Phạm vi xuất báo cáo",
            ["Tất cả các khung đánh giá", "Chọn khung đánh giá cụ thể"]
        )
        
        filtered_data = company_data
        filtered_participants = company_participants
        
        if frame_option == "Chọn khung đánh giá cụ thể":
            selected_frame = st.selectbox("Chọn khung đánh giá", options=frames, key="export_frame")
            filtered_data = company_data[company_data['frame_id'] == selected_frame]
            filtered_participants = company_participants[company_participants['frame_id'] == selected_frame]
        
        # Prepare data for export
        audit_data = []
        for _, row in filtered_data.iterrows():
            audit_data.append({
                'company': row['company'],
                'address': row['address'],
                'department': row['department'],
                'person': row['person'],
                'audit_time': row['audit_time'],
                'frame_id': row['frame_id'],
                'panel_id': row['panel_id'],
                'clause': row['clause'],
                'clause_name': row['clause_name'],
                'requirements': row['requirements'],
                'evidence': row['evidence'],
                'image_url': row['image_url'],
                'result': row['result'],
                'auditor': row['auditor'],
                'timestamp': row['timestamp']
            })
        
        participants_data = []
        for _, row in filtered_participants.iterrows():
            participants_data.append({
                'company': row['company'],
                'frame_id': row['frame_id'],
                'fullname': row['fullname'],
                'position': row['position'],
                'role': row['role']
            })
        
        # Add export buttons
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("Xuất PDF"):
                with st.spinner("Đang tạo file PDF..."):
                    pdf_data = export_to_pdf(selected_company, audit_data, participants_data)
                    company_name_safe = selected_company.replace(' ', '_')
                    date_str = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"bao_cao_danh_gia_iso_{company_name_safe}_{date_str}.pdf"
                    
                    st.markdown(
                        get_download_link(pdf_data, filename, "📥 Tải xuống file PDF"),
                        unsafe_allow_html=True
                    )
        
        with col2:
            if st.button("Xuất Word"):
                with st.spinner("Đang tạo file Word..."):
                    docx_data = export_to_word(selected_company, audit_data, participants_data)
                    company_name_safe = selected_company.replace(' ', '_')
                    date_str = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"bao_cao_danh_gia_iso_{company_name_safe}_{date_str}.docx"
                    
                    st.markdown(
                        get_download_link(docx_data, filename, "📥 Tải xuống file Word"),
                        unsafe_allow_html=True
                    )
    else:
        st.warning("Không có khung đánh giá nào cho công ty này.")

# ============ Main App ============
def main():
    # Load CSS
    load_css()
    
    # Initialize session state variables if they don't exist
    if "is_logged_in" not in st.session_state:
        st.session_state.is_logged_in = False
    
    if "page" not in st.session_state:
        st.session_state.page = "main"
    
    if "show_register" not in st.session_state:
        st.session_state.show_register = False
    
    # Router based on login status and current page
    if not st.session_state.is_logged_in:
        page_login()
    else:
        if st.session_state.page == "change_password":
            page_change_password()
        else:
            page_main()

if __name__ == "__main__":
    main()
